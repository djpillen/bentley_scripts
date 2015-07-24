from lxml import etree
import csv
from docx import Document
import re
from os.path import join

path = 'C:/Users/djpillen/GitHub/vandura/Real_Masters_all'

#title_file = 'collection_titles.csv'
top_level_file = 'C:/Users/Public/Documents/accessrestrict_toplevel-4.csv'
expired_file = 'C:/Users/Public/Documents/accessrestrict_expired_fulltext-1.csv'
all_dates_file = 'C:/Users/Public/Documents/accessrestrictdate-8.csv'
no_date_file = 'C:/Users/Public/Documents/accessrestrict_nodate-1.csv'

access_restrictions = {}
current_dict = {}
non_time = {}

with open(expired_file, 'rb') as expired_csv:
    reader = csv.reader(expired_csv)
    next(reader, None)
    for row in reader:
        filename = row[0]
        print 'Parsing expired_file for', filename
        date_path = row[1]
        expiration_date = row[2]
        full_text = row[3]
        full_text = re.sub(r'</?(.*?)>','', full_text)
        full_text = full_text
        if filename not in access_restrictions:
            access_restrictions[filename] = {}
        if 'Expired' not in access_restrictions[filename]:
            access_restrictions[filename]['Expired'] = {}

        access_restrictions[filename]['Expired'][date_path] = {}
        access_restrictions[filename]['Expired'][date_path]['Expiration_date'] = expiration_date
        access_restrictions[filename]['Expired'][date_path]['Full_text'] = full_text

        component_path = re.sub(r'/accessrestrict/p/date','',date_path)
        component_path = component_path
        tree = etree.parse(join(path,filename))
        for r in tree.xpath(component_path):
            container = r.xpath('.//container')
            physloc = r.xpath('.//physloc')
            if container:
                if container[0].text:
                    box_number = container[0].text
                else:
                    box_number = 'N/A'
            elif physloc:
                if physloc[0].text:
                    box_number = physloc[0].text
                else:
                    box_number = 'N/A'
            else:
                box_number = 'N/A'
            access_restrictions[filename]['Expired'][date_path]['Container_number'] = box_number

"""
with open(title_file,'rb')as titles_csv:
    reader = csv.reader(titles_csv)
    next(reader, None)
    for row in reader:
        filename = row[0]
        coll_title = row[1]
        if filename in access_restrictions:
            access_restrictions[filename]['Title'] = coll_title
"""

with open(top_level_file, 'rb') as top_csv:
    reader = csv.reader(top_csv)
    for row in reader:
        filename = row[0]
        print 'Parsing top_level_file for', filename
        restriction = row[2]
        restriction = restriction.replace('<item>','\n*')
        restriction = re.sub(r'</?(.*?)>','', restriction)
        restriction = restriction
        if filename in access_restrictions:
            access_restrictions[filename]['Restriction'] = restriction

with open(all_dates_file,'rb') as all_csv:
    reader = csv.reader(all_csv)
    for row in reader:
        filename = row[0]
        print 'Parsing all_date_file for', filename
        normal = row[3]
        if normal > '2015-07-02':
            if filename not in current_dict:
                current_dict[filename] = 1
            else:
                current_dict[filename] += 1


with open(no_date_file, 'rb') as no_date_csv:
    reader = csv.reader(no_date_csv)
    for row in reader:
        filename = row[0]
        print 'Parsing no_date_file for', filename
        if filename not in non_time:
            non_time[filename] = 1
        else:
            non_time[filename] += 1


print 'Parsing dictionaries and building document'

for filename in access_restrictions:
    if filename in current_dict:
        access_restrictions[filename]['Remaining_time'] = str(current_dict[filename])
    else:
        access_restrictions[filename]['Remaining_time'] = '0'

    if filename in non_time:
        access_restrictions[filename]['Remaining_non_time'] = str(non_time[filename])
    else:
        access_restrictions[filename]['Remaining_non_time'] = '0'


document = Document()

for filename in access_restrictions:
    document.add_heading(filename)
    document.add_heading('Collection Level Restriction', level=2)
    document.add_paragraph(access_restrictions[filename]['Restriction'])
    document.add_heading('Expired Restrictions', level=2)
    table = document.add_table(rows=1, cols=3)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Container"
    hdr_cells[1].text = 'Expiration Date'
    hdr_cells[2].text = 'Full Restriction'
    for date_path in access_restrictions[filename]['Expired']:
        row_cells = table.add_row().cells
        row_cells[0].text = access_restrictions[filename]['Expired'][date_path]['Container_number']
        row_cells[1].text = access_restrictions[filename]['Expired'][date_path]['Expiration_date']
        row_cells[2].text = access_restrictions[filename]['Expired'][date_path]['Full_text']


    table.style= "TableGrid"
    document.add_paragraph('Remaining Time-Bound Restrictions')
    document.add_paragraph(access_restrictions[filename]['Remaining_time'])
    document.add_paragraph('Remaining Non-Time-Bound Restrictions')
    document.add_paragraph(access_restrictions[filename]['Remaining_non_time'])
    document.add_page_break()

document.save('C:/Users/Public/Documents/expired_restrictions_20150713.docx')
