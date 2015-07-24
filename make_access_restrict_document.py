from lxml import etree
import csv
from docx import Document
import re
from os.path import join

path = 'C:/Users/djpillen/GitHub/vandura/Real_Masters_all'

#title_file = 'collection_titles.csv'
top_level_file = 'C:/Users/Public/Documents/accessrestrict_toplevel-5.csv'
component_file = 'C:/Users/Public/Documents/accessrestrict_component-6.csv'

access_restrictions = {}

with open(component_file, 'rb') as component_csv:
    reader = csv.reader(component_csv)
    next(reader, None)
    for row in reader:
        filename = row[0]
        print 'Parsing component_file for', filename
        access_path = row[1]
        if filename not in access_restrictions:
            access_restrictions[filename] = {}
        if 'Components' not in access_restrictions[filename]:
            access_restrictions[filename]['Components'] = 1
        else:
            access_restrictions[filename]['Components'] += 1
        tree = etree.parse(join(path,filename))
        accessrestrict = tree.xpath(access_path)
        component = accessrestrict[0].getparent()
        container = component.xpath('./did/container')
        physloc = component.xpath('./did/physloc')
        if container:
            if container[0].text:
                box_number = container[0].text
            else:
                box_number = 'N/A'
        elif physloc:
            if physloc[0].text:
                box_number = physloc[0].text
            else:
                box_number = 'N/A'
        elif not container and not physloc:
            box_number = 'N/A'
        print box_number
        if 'Containers' not in access_restrictions[filename]:
            access_restrictions[filename]['Containers'] = []
        if box_number not in access_restrictions[filename]['Containers']:
            access_restrictions[filename]['Containers'].append(box_number)


with open(top_level_file, 'rb') as top_csv:
    reader = csv.reader(top_csv)
    for row in reader:
        filename = row[0]
        print 'Parsing top_level_file for', filename
        restriction = row[2]
        restriction = restriction.replace('<item>','\n*')
        restriction = re.sub(r'</?(.*?)>','', restriction)
        restriction = restriction
        if filename not in access_restrictions:
            access_restrictions[filename] = {}
            access_restrictions[filename]['Restriction'] = restriction
        elif filename in access_restrictions:
            access_restrictions[filename]['Restriction'] = restriction

for filename in access_restrictions:
    if 'Containers' not in access_restrictions[filename]:
        access_restrictions[filename]['Containers'] = []
        access_restrictions[filename]['Containers'].append("No containers")
    if 'Components' not in access_restrictions[filename]:
        access_restrictions[filename]['Components'] = 0


print 'Parsing dictionaries and building document'


document = Document()

for filename in sorted(access_restrictions):
    document.add_heading(filename)
    document.add_heading('Collection Level Restriction', level=2)
    document.add_paragraph(access_restrictions[filename]['Restriction'])
    document.add_heading('Component Level Restrictions', level=2)
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Component Restrictions"
    hdr_cells[1].text = 'Containers'
    row_cells = table.add_row().cells
    row_cells[0].text = str(access_restrictions[filename]['Components'])
    row_cells[1].text = ', '.join(item for item in sorted(access_restrictions[filename]['Containers']))
    document.add_page_break()


document.save('C:/Users/Public/Documents/restrictions.docx')
